"""Tests for Word document loader — including math formulas and text boxes.

.. note::

    .docx files are ZIP archives containing ``word/document.xml``.
    Fixtures that inject formulas / text boxes must modify the XML
    inside the ZIP rather than calling ``etree.parse()`` on the file directly.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from lxml import etree
from src.domain import Document, ProcessingStatus
from src.parser.loader.word_loader import WordLoader

# OOXML namespaces
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"


# ---------------------------------------------------------------------------
# Helpers for editing word/document.xml inside a .docx ZIP
# ---------------------------------------------------------------------------


def _read_docx_xml(filepath: str) -> etree._Element:
    """Read and parse ``word/document.xml`` from a .docx ZIP."""
    with zipfile.ZipFile(filepath, "r") as zf:
        xml_bytes = zf.read("word/document.xml")
    return etree.fromstring(xml_bytes)


def _write_docx_xml(filepath: str, root: etree._Element) -> None:
    """Replace ``word/document.xml`` inside the .docx ZIP."""
    from tempfile import NamedTemporaryFile

    xml_bytes = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    # Read all entries into memory first so we can close the input ZIP
    # before opening the output (avoids Windows file-lock issues).
    with zipfile.ZipFile(filepath, "r") as zin:
        entries = [(item, zin.read(item.filename)) for item in zin.infolist()]

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmpname = tmp.name
        with zipfile.ZipFile(tmpname, "w") as zout:
            for item, data in entries:
                if item.filename == "word/document.xml":
                    zout.writestr(item, xml_bytes)
                else:
                    zout.writestr(item, data)

    Path(tmpname).replace(filepath)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def sample_docx(tmp_path: Path) -> str:
    """Create a minimal .docx with headings, paragraphs, and a table."""
    from docx import Document

    doc = Document()
    doc.add_heading("第一章 引言", level=1)
    doc.add_paragraph("这是引言段落的内容。")
    doc.add_paragraph("第二段正文，继续描述背景信息。")
    doc.add_heading("1.1 研究背景", level=2)
    doc.add_paragraph("研究背景的具体描述内容。")

    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "数量"
    table.cell(0, 2).text = "备注"
    table.cell(1, 0).text = "文档"
    table.cell(1, 1).text = "10"
    table.cell(1, 2).text = "已完成"

    fp = tmp_path / "test.docx"
    doc.save(str(fp))
    return str(fp)


@pytest.fixture
def docx_with_formula(tmp_path: Path) -> str:
    """Create a .docx where one paragraph contains an OMML math formula.

    python-docx cannot add math natively, so we inject the ``m:oMath`` XML
    directly into ``word/document.xml`` inside the ZIP.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("本页包含一个公式：")
    doc.add_paragraph("")  # placeholder — formula injected here
    doc.add_paragraph("以上就是公式内容。")

    fp = tmp_path / "formula_test.docx"
    doc.save(str(fp))

    # ── Inject OMML math into paragraph 1 via XML ──
    root = _read_docx_xml(str(fp))
    body = root.find(f".//{{{_W_NS}}}body")
    assert body is not None, "No <w:body> found"
    paras = body.findall(f".//{{{_W_NS}}}p")
    assert len(paras) >= 2, "Expected at least 2 paragraphs"

    target_p = paras[1]  # second paragraph (0-indexed)

    omath = etree.SubElement(target_p, f"{{{_M_NS}}}oMath")
    mr = etree.SubElement(omath, f"{{{_M_NS}}}r")
    mt = etree.SubElement(mr, f"{{{_M_NS}}}t")
    mt.text = "x = (-b ± √(b²-4ac)) / 2a"

    _write_docx_xml(str(fp), root)
    return str(fp)


@pytest.fixture
def docx_with_textbox(tmp_path: Path) -> str:
    """Create a .docx with a floating text box (``w:txbxContent``).

    We inject a ``<wps:wsp>`` containing ``<wps:txbx><w:txbxContent>…``
    into the document XML.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("这是文档正文。")
    doc.add_paragraph("正文第二行。")

    fp = tmp_path / "textbox_test.docx"
    doc.save(str(fp))

    # ── Inject a text box via XML ──
    root = _read_docx_xml(str(fp))
    body = root.find(f".//{{{_W_NS}}}body")
    assert body is not None

    # Create a new paragraph with the text box drawing
    p = etree.SubElement(body, f"{{{_W_NS}}}p")
    r1 = etree.SubElement(p, f"{{{_W_NS}}}r")

    # mc:AlternateContent
    alt = etree.SubElement(r1, f"{{{_MC_NS}}}AlternateContent")
    choice = etree.SubElement(alt, f"{{{_MC_NS}}}Choice")
    choice.set("Requires", "wps")

    drawing = etree.SubElement(choice, f"{{{_WP_NS}}}inline")
    _add_extent(drawing, "914400", "914400")
    effect_extent = etree.SubElement(drawing, f"{{{_WP_NS}}}effectExtent")
    for a, v in [("l", "0"), ("t", "0"), ("r", "0"), ("b", "0")]:
        effect_extent.set(a, v)

    doc_pr = etree.SubElement(drawing, f"{{{_WP_NS}}}docPr")
    doc_pr.set("id", "1")
    doc_pr.set("name", "Text Shape 1")

    graphic = etree.SubElement(drawing, f"{{{_A_NS}}}graphic")
    graphic_data = etree.SubElement(graphic, f"{{{_A_NS}}}graphicData")
    graphic_data.set(
        "uri", "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    )

    wsp = etree.SubElement(graphic_data, f"{{{_WPS_NS}}}wsp")
    txbx = etree.SubElement(wsp, f"{{{_WPS_NS}}}txbx")
    txbx_content = etree.SubElement(txbx, f"{{{_W_NS}}}txbxContent")

    tbp = etree.SubElement(txbx_content, f"{{{_W_NS}}}p")
    tbr = etree.SubElement(tbp, f"{{{_W_NS}}}r")
    tbt = etree.SubElement(tbr, f"{{{_W_NS}}}t")
    tbt.text = "这是悬浮文本框中的内容"
    tbt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    # Fallback for the AlternateContent
    fallback = etree.SubElement(alt, f"{{{_MC_NS}}}Fallback")
    fb_r = etree.SubElement(fallback, f"{{{_W_NS}}}r")
    fb_t = etree.SubElement(fb_r, f"{{{_W_NS}}}t")
    fb_t.text = "[文本框]"
    fb_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    _write_docx_xml(str(fp), root)
    return str(fp)


def _add_extent(parent, cx: str, cy: str):
    """Add an ``a:extent`` child with cx/cy attributes."""
    ext = etree.SubElement(parent, f"{{{_A_NS}}}extent")
    ext.set("cx", cx)
    ext.set("cy", cy)


# ---------------------------------------------------------------------------
# Basic tests
# ---------------------------------------------------------------------------


class TestWordLoader:
    """Basic document structure tests."""

    def test_load_returns_document(self, sample_docx: str):
        loader = WordLoader()
        doc = loader.load(sample_docx)
        assert isinstance(doc, Document)
        assert doc.filename == "test.docx"
        assert doc.file_type == "docx"
        assert doc.status == ProcessingStatus.UPLOADED

    def test_load_single_page(self, sample_docx: str):
        loader = WordLoader()
        doc = loader.load(sample_docx)
        assert len(doc.pages) == 1
        assert doc.pages[0].page_num == 1

    def test_load_blocks_with_headings(self, sample_docx: str):
        loader = WordLoader()
        doc = loader.load(sample_docx)
        blocks = doc.pages[0].blocks
        headings = [b for b in blocks if b.block_type in ("title", "section_heading")]
        assert len(headings) >= 2

        title_block = headings[0]
        assert "第一章" in title_block.content
        assert title_block.block_type == "title"

        sub_heading = headings[1]
        assert "研究背景" in sub_heading.content
        assert sub_heading.block_type == "section_heading"

    def test_load_table_block(self, sample_docx: str):
        loader = WordLoader()
        doc = loader.load(sample_docx)
        blocks = doc.pages[0].blocks
        tables = [b for b in blocks if b.block_type == "table"]
        assert len(tables) == 1
        assert "项目" in tables[0].content
        assert "文档" in tables[0].content

    def test_load_accumulates_text(self, sample_docx: str):
        loader = WordLoader()
        doc = loader.load(sample_docx)
        page_text = doc.pages[0].text
        assert "第一章 引言" in page_text
        assert "研究背景" in page_text
        assert "项目" in page_text

    def test_load_empty_paragraph_skipped(self, tmp_path: Path):
        from docx import Document

        docx = Document()
        docx.add_paragraph("")
        docx.add_paragraph("real content")
        fp = tmp_path / "empty_test.docx"
        docx.save(str(fp))

        loader = WordLoader()
        doc = loader.load(str(fp))
        blocks = doc.pages[0].blocks
        assert all(b.content.strip() for b in blocks)
        assert len(blocks) == 1

    def test_load_nonexistent_file(self):
        loader = WordLoader()
        with pytest.raises(FileNotFoundError):
            loader.load("/nonexistent/path.docx")

    def test_load_unsupported_format(self, tmp_path: Path):
        loader = WordLoader()
        f = tmp_path / "test.txt"
        f.write_text("not a docx")
        with pytest.raises(ValueError, match="不支持的文件格式"):
            loader.load(str(f))

    def test_reading_order_preserved(self, sample_docx: str):
        loader = WordLoader()
        doc = loader.load(sample_docx)
        blocks = doc.pages[0].blocks
        orders = [b.reading_order for b in blocks]
        assert orders == sorted(orders)

    def test_metadata_on_blocks(self, sample_docx: str):
        loader = WordLoader()
        doc = loader.load(sample_docx)
        text_blocks = [b for b in doc.pages[0].blocks if b.block_type == "text"]
        if text_blocks:
            assert "style" in text_blocks[0].metadata

    def test_word_doc_extension(self, tmp_path: Path):
        from docx import Document

        docx = Document()
        docx.add_paragraph("legacy test")
        fp = tmp_path / "legacy.doc"
        docx.save(str(fp))
        loader = WordLoader()
        doc = loader.load(str(fp))
        assert doc.file_type == "docx"
        assert len(doc.pages[0].blocks) == 1


# ---------------------------------------------------------------------------
# Math formula tests
# ---------------------------------------------------------------------------


class TestWordLoaderMath:
    """OMML math formula extraction."""

    def test_detects_math_formula(self, docx_with_formula: str):
        loader = WordLoader()
        doc = loader.load(docx_with_formula)
        formula_blocks = [b for b in doc.pages[0].blocks if b.block_type == "formula"]
        assert len(formula_blocks) >= 1

    def test_math_formula_content(self, docx_with_formula: str):
        loader = WordLoader()
        doc = loader.load(docx_with_formula)
        for block in doc.pages[0].blocks:
            if block.block_type == "formula":
                assert "x =" in block.content
                assert "2a" in block.content
                return
        pytest.fail("No formula block found")

    def test_math_block_metadata(self, docx_with_formula: str):
        loader = WordLoader()
        doc = loader.load(docx_with_formula)
        for block in doc.pages[0].blocks:
            if block.block_type == "formula":
                assert block.metadata.get("has_formula") is True
                return
        pytest.fail("No formula block found")

    def test_non_math_paragraphs_unaffected(self, docx_with_formula: str):
        loader = WordLoader()
        doc = loader.load(docx_with_formula)
        texts = [b.content for b in doc.pages[0].blocks if b.block_type == "text"]
        assert any("公式" in t for t in texts)
        assert any("以上" in t for t in texts)


# ---------------------------------------------------------------------------
# Text box tests
# ---------------------------------------------------------------------------


class TestWordLoaderTextbox:
    """Floating text box / shape extraction."""

    def test_textbox_content_captured(self, docx_with_textbox: str):
        loader = WordLoader()
        doc = loader.load(docx_with_textbox)
        found = any("文本框" in b.content for b in doc.pages[0].blocks)
        assert found, "Text box content was not extracted"

    def test_textbox_source_metadata(self, docx_with_textbox: str):
        loader = WordLoader()
        doc = loader.load(docx_with_textbox)
        textbox_blocks = [
            b
            for b in doc.pages[0].blocks
            if b.metadata.get("source") == "textbox"
        ]
        assert len(textbox_blocks) >= 1

    def test_textbox_doesnt_displace_normal_paragraphs(self, docx_with_textbox: str):
        loader = WordLoader()
        doc = loader.load(docx_with_textbox)
        normal = [
            b
            for b in doc.pages[0].blocks
            if b.block_type == "text" and b.metadata.get("source") != "textbox"
        ]
        assert len(normal) >= 2
