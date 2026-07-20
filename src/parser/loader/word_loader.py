"""Word document loader using python-docx + lxml XML traversal.

Extracts:
  - Paragraphs with heading-level detection (style + font-size heuristics)
  - **Math formulas** (OMML ``m:oMath``) — missing from python-docx ``paragraph.text``
  - **Text boxes / floating shapes** (``w:txbxContent``) — invisible to the high-level API
  - Tables as structured ``Block`` objects

Since .docx has no physical page concept, the entire document is treated as a
single logical page for pipeline compatibility.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument

from .base import DocumentLoader
from src.domain import Document, Page, Block
from src.domain.enums import ProcessingStatus

# ---------------------------------------------------------------------------
# OOXML namespace constants
# ---------------------------------------------------------------------------
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"

_NSMAP = {
    "w": _W_NS,
    "m": _M_NS,
    "r": _R_NS,
    "wp": _WP_NS,
    "a": _A_NS,
    "pic": _PIC_NS,
}

# ---------------------------------------------------------------------------
# Heading detection
# ---------------------------------------------------------------------------
_HEADING_LEVEL_MAP = {
    "Heading 1": "title",
    "Heading 2": "section_heading",
    "Heading 3": "section_heading",
    "Heading 4": "section_heading",
    "Heading 5": "section_heading",
    "Heading 6": "section_heading",
    "Title": "title",
    "Subtitle": "section_heading",
    "1": "title",
    "2": "section_heading",
    "3": "section_heading",
}

_CHARS_PER_PAGE_ESTIMATE = 3000
_DEFAULT_PAGE_WIDTH = 595  # A4 pt
_DEFAULT_PAGE_HEIGHT = 842  # A4 pt


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------


def _collect_text_from_element(element) -> str:
    """Walk *element* in document order collecting text from ``w:t`` and ``m:t``.

    This captures both regular paragraph text and OMML math formula text,
    which python-docx's ``paragraph.text`` silently drops.
    """
    parts: list[str] = []
    for child in element.iter():
        tag = child.tag
        if "}" not in tag:
            continue
        ns = tag[tag.index("{") + 1 : tag.index("}")]
        local = tag[tag.rindex("}") + 1 :]
        if local == "t" and ns in (_W_NS, _M_NS) and child.text:
            parts.append(child.text)
    return "".join(parts)


def _has_math(element) -> bool:
    """Return ``True`` if *element* contains OMML math formula elements."""
    return len(element.findall(".//m:oMath", _NSMAP)) > 0


def _extract_math_text(element) -> str:
    """Extract the linear plain-text form of all math formulas in *element*.

    Used as fallback when LaTeX conversion produces empty output.
    """
    parts: list[str] = []
    for math_elem in element.findall(".//m:oMath", _NSMAP):
        formula_text = _collect_text_from_element(math_elem)
        if formula_text.strip():
            parts.append(formula_text.strip())
    return " ".join(parts)


# ---------------------------------------------------------------------------
# OMML → LaTeX converter
# ---------------------------------------------------------------------------

_M_TAG = f"{{{_M_NS}}}"  # e.g. "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def _ltag(el) -> str:
    """Extract the local (namespace-stripped) tag name from an lxml element."""
    tag = el.tag
    return tag[tag.rindex("}") + 1:] if "}" in tag else tag


def _omml_to_latex(element) -> str:
    """Recursively convert an OMML element subtree to a LaTeX string.

    Handles the most common OMML constructs (fractions, superscripts,
    subscripts, radicals, n-ary operators, accents, delimiters) and
    falls back to plain text for unhandled elements.
    """
    local = _ltag(element)

    # --- 文本叶节点 ---
    if local == "r":
        for child in element:
            if _ltag(child) == "t" and child.text:
                return _map_math_text(child.text)
        return ""
    if local == "t":
        return _map_math_text(element.text or "")

    # --- 属性 / 控制符 → 跳过 ---
    if local in ("spPr", "ctrl", "rPr", "eProps", "ctlPr"):
        return ""

    # --- 分数 ---
    if local == "f":
        num = ""
        den = ""
        for child in element:
            cl = _ltag(child)
            if cl == "num":
                num = _omml_to_latex(child)
            elif cl == "den":
                den = _omml_to_latex(child)
        return _latex_cmd(r"\frac", num, den)

    # --- 上标 ---
    if local == "sup":
        parts = [_omml_to_latex(c) for c in element]
        if len(parts) >= 2:
            return f"{{{' '.join(parts[:-1])}}}^{{{parts[-1]}}}"
        return "".join(parts)
    if local == "sSup":  # 无基底上标
        sup = ""
        for child in element:
            cl = _ltag(child)
            if cl == "e":
                sup = _omml_to_latex(child)
        return f"^{{{sup}}}" if sup else ""

    # --- 下标 ---
    if local == "sub":
        parts = [_omml_to_latex(c) for c in element]
        if len(parts) >= 2:
            return f"{{{' '.join(parts[:-1])}}}_{{{parts[-1]}}}"
        return "".join(parts)
    if local == "sSub":  # 无基底下标
        sub = ""
        for child in element:
            cl = _ltag(child)
            if cl == "e":
                sub = _omml_to_latex(child)
        return f"_{{{sub}}}" if sub else ""

    # --- 上下标组合 (m:pre = presubscript/presuperscript) ---
    if local == "pre":
        sub = sup = base = ""
        for child in element:
            cl = _ltag(child)
            if cl == "sub":
                sub = _omml_to_latex(child)
            elif cl == "sup":
                sup = _omml_to_latex(child)
            elif cl == "e":
                base = _omml_to_latex(child)
        return f"{{{base}}}^{{{sup}}}_{{{sub}}}" if (sub or sup) else base

    # --- 根号 ---
    if local == "rad":
        rad = ""
        deg = ""
        for child in element:
            cl = _ltag(child)
            if cl == "e":
                rad = _omml_to_latex(child)
            elif cl == "deg":
                deg = _omml_to_latex(child)
        if deg:
            return _latex_cmd(r"\sqrt", _maybe_brace(deg), rad)
        return _latex_cmd(r"\sqrt", rad)

    # --- N 元运算符 (求和、积分、乘积) ---
    if local == "nary":
        op = sub = sup = base = ""
        for child in element:
            cl = _ltag(child)
            if cl == "naryPr":
                for prop in child:
                    pl = _ltag(prop)
                    if pl == "chr" and prop.text:
                        op = _OMML_NARY.get(prop.text, prop.text)
            elif cl == "sub":
                sub = _omml_to_latex(child)
            elif cl == "sup":
                sup = _omml_to_latex(child)
            elif cl == "e":
                base = _omml_to_latex(child)
        _op = op or "\\sum"
        result = _op
        if sub:
            result += f"_{{{sub}}}"
        if sup:
            result += f"^{{{sup}}}"
        if base:
            result += f" {{{base}}}"
        return result

    # --- 着重号 (acc) ---
    if local == "acc":
        base = ""
        accent = "̂"  # U+0302 combining circumflex
        for child in element:
            cl = _ltag(child)
            if cl == "accPr":
                for prop in child:
                    pl = _ltag(prop)
                    if pl == "chr" and prop.text:
                        accent = _OMML_ACCENT.get(prop.text, prop.text)
            elif cl == "e":
                base = _omml_to_latex(child)
        # Map combining char to LaTeX command
        latex_accent = _ACCENT_MAP.get(accent, f"\\hat")
        return _latex_cmd(latex_accent, base)

    # --- 上划线/下划线 ---
    if local == "bar":
        base = ""
        pos = "top"
        for child in element:
            cl = _ltag(child)
            if cl == "barPr":
                for prop in child:
                    pl = _ltag(prop)
                    if pl == "pos" and prop.text:
                        pos = prop.text
            elif cl == "e":
                base = _omml_to_latex(child)
        cmd = r"\overline" if pos in ("top", "topBot") else r"\underline"
        return _latex_cmd(cmd, base)

    # --- 定界符 (括号) ---
    if local == "d":
        sep = ""
        inner = ""
        for child in element:
            cl = _ltag(child)
            if cl == "dPr":
                for prop in child:
                    pl = _ltag(prop)
                    if pl == "sepChr" and prop.text:
                        sep = prop.text
            elif cl == "e":
                inner = _omml_to_latex(child)
        if sep:
            # 矩阵等用分隔符的情况 → left/right + sep
            return f"\\left({inner}\\right)"
        return f"\\left({inner}\\right)"  # default parentheses

    # --- 分组 / 盒子 → 递归子节点 ---
    if local in ("e", "box", "group", "func"):
        return _omml_children(element)

    # --- 极限 (limLow / limUpp) ---
    if local == "limLow":
        base = ""
        limit = ""
        for child in element:
            cl = _ltag(child)
            if cl == "e":
                if not base:
                    base = _omml_to_latex(child)
                else:
                    limit = _omml_to_latex(child)
            elif cl == "lim":
                limit = _omml_to_latex(child)
        if limit:
            return f"{{{base}}}_{{{limit}}}"
        return base
    if local == "limUpp":
        base = ""
        limit = ""
        for child in element:
            cl = _ltag(child)
            if cl == "e":
                if not base:
                    base = _omml_to_latex(child)
                else:
                    limit = _omml_to_latex(child)
            elif cl == "lim":
                limit = _omml_to_latex(child)
        if limit:
            return f"{{{base}}}^{{{limit}}}"
        return base

    # --- 矩阵 ---
    if local == "m":
        rows = []
        current_row: list[str] = []
        for child in element:
            cl = _ltag(child)
            if cl == "mr":
                if current_row:
                    rows.append(" & ".join(current_row))
                    current_row = []
            elif cl == "e":
                current_row.append(_omml_to_latex(child))
        if current_row:
            rows.append(" & ".join(current_row))
        col_align = "c" * (max((r.count(" & ") + 1 for r in rows), default=1))
        body = r"\\".join(rows)
        return f"\\begin{{matrix}} {body} \\end{{matrix}}"

    # --- 方程数组 ---
    if local == "eqArr":
        rows = []
        current_row: list[str] = []
        for child in element:
            cl = _ltag(child)
            if cl == "e":
                current_row.append(_omml_to_latex(child))
            elif cl == "eqArrPr":
                pass
        if current_row:
            rows.append(" & ".join(current_row))
        body = r"\\".join(rows)
        return f"\\begin{{aligned}} {body} \\end{{aligned}}"

    # --- 缺省：递归子节点 ---
    return _omml_children(element)


def _omml_children(element) -> str:
    """Concatenate LaTeX output of all direct children (skipping properties)."""
    parts = []
    for child in element:
        if _ltag(child) not in ("spPr", "rPr", "eProps", "ctlPr"):
            parts.append(_omml_to_latex(child))
    return "".join(parts)


def _map_math_text(text: str) -> str:
    """Map Unicode math symbols and Greek letter names in *text*."""
    # Greek letters stored as named entities by the OMML authoring tool
    result = _OMML_GREEK_NAMES.get(text, text)
    # Unicode math symbols
    result = _OMML_SYMBOLS.get(result, result)
    return result


def _latex_cmd(cmd: str, *args: str) -> str:
    """Build a LaTeX command: ``\\cmd{arg1}{arg2}``."""
    return cmd + "".join(f"{{{a}}}" for a in args)


def _maybe_brace(s: str) -> str:
    """Wrap *s* in braces if it contains special chars or is longer than 1 char."""
    if not s:
        return "{}"
    if len(s) == 1 and s.isalnum():
        return s
    return f"{{{s}}}"


# OMML n-ary operator characters
_OMML_NARY = {
    "∑": r"\sum",
    "∏": r"\prod",
    "∫": r"\int",
    "∐": r"\coprod",
    "⋃": r"\bigcup",
    "⋂": r"\bigcap",
    "⊕": r"\bigoplus",
    "⊗": r"\bigotimes",
}

# OMML accent characters → LaTeX command
_OMML_ACCENT = {
    "̀":  r"\grave",
    "́":  r"\acute",
    "̂":  r"\hat",
    "̃":  r"\tilde",
    "̄":  r"\bar",
    "̆":  r"\breve",
    "̇":  r"\dot",
    "̈":  r"\ddot",
    "⃗":  r"\vec",
}
_ACCENT_MAP = {v: _OMML_ACCENT.get(k, r"\hat") for k, v in _OMML_ACCENT.items()}
# Also map by combining character directly
_ACCENT_MAP.update(_OMML_ACCENT)

# OMML Greek letter names (common subset)
_OMML_GREEK_NAMES = {
    "alpha": r"\alpha", "beta": r"\beta", "gamma": r"\gamma",
    "delta": r"\delta", "epsilon": r"\epsilon", "zeta": r"\zeta",
    "eta": r"\eta", "theta": r"\theta", "iota": r"\iota",
    "kappa": r"\kappa", "lambda": r"\lambda", "mu": r"\mu",
    "nu": r"\nu", "xi": r"\xi", "omicron": r"o",
    "pi": r"\pi", "rho": r"\rho", "sigma": r"\sigma",
    "tau": r"\tau", "upsilon": r"\upsilon", "phi": r"\phi",
    "chi": r"\chi", "psi": r"\psi", "omega": r"\omega",
    "Alpha": r"\Alpha", "Beta": r"\Beta", "Gamma": r"\Gamma",
    "Delta": r"\Delta", "Epsilon": r"\Epsilon", "Zeta": r"\Zeta",
    "Eta": r"\Eta", "Theta": r"\Theta", "Iota": r"\Iota",
    "Kappa": r"\Kappa", "Lambda": r"\Lambda", "Mu": r"\Mu",
    "Nu": r"\Nu", "Xi": r"\Xi", "Omicron": r"O",
    "Pi": r"\Pi", "Rho": r"\Rho", "Sigma": r"\Sigma",
    "Tau": r"\Tau", "Upsilon": r"\Upsilon", "Phi": r"\Phi",
    "Chi": r"\Chi", "Psi": r"\Psi", "Omega": r"\Omega",
}

# Common Unicode math symbols → LaTeX
_OMML_SYMBOLS = {
    "∞": r"\infty", "∂": r"\partial", "∇": r"\nabla",
    "→": r"\to", "⇒": r"\Rightarrow", "↦": r"\mapsto",
    "≈": r"\approx", "≃": r"\simeq", "≅": r"\cong",
    "≠": r"\neq", "≡": r"\equiv", "≤": r"\leq", "≥": r"\geq",
    "∈": r"\in", "∉": r"\notin", "∋": r"\ni",
    "⊂": r"\subset", "⊃": r"\supset", "⊆": r"\subseteq", "⊇": r"\supseteq",
    "∪": r"\cup", "∩": r"\cap", "∖": r"\setminus",
    "∀": r"\forall", "∃": r"\exists", "∄": r"\nexists",
    "∧": r"\land", "∨": r"\lor", "¬": r"\lnot",
    "⊥": r"\perp", "∥": r"\parallel",
    "∠": r"\angle", "△": r"\triangle",
    "⊕": r"\oplus", "⊗": r"\otimes", "⊙": r"\odot",
    "±": r"\pm", "∓": r"\mp", "×": r"\times", "÷": r"\div",
    "⋅": r"\cdot", "∘": r"\circ", "∗": r"\ast",
    "√": r"\sqrt",
    "∛": r"\sqrt[3]",
    "∜": r"\sqrt[4]",
}


def _extract_math_latex(element) -> str:
    """Convert all OMML formulas in *element* to LaTeX.

    Returns formulas separated by newlines, each wrapped in ``$$...$$``
    for clear identification in the chunk content.
    """
    parts: list[str] = []
    for math_elem in element.findall(".//m:oMath", _NSMAP):
        latex = _omml_to_latex(math_elem)
        if latex.strip():
            parts.append(f"$${latex.strip()}$$")
    return "\n".join(parts)


def _normalize_image_bytes(blob: bytes) -> bytes:
    """Normalize embedded image bytes to a raster format (PNG).

    Word documents produced by Equation Editor / MathType store formula
    images as **WMF/EMF vector metafiles**, which neither the OCR pipeline
    nor the Qwen-VL API can consume directly (they expect PNG/JPEG bytes).
    This detects metafiles via Pillow and rasterizes them to PNG at a
    readable DPI, falling back to the original bytes on any failure.
    """
    try:
        import io
        from PIL import Image

        im = Image.open(io.BytesIO(blob))
        if im.format not in ("WMF", "EMF"):
            return blob  # already a raster format — leave untouched

        # Render the vector metafile at high DPI for legible formulas
        try:
            im.load(dpi=300)
        except Exception:
            pass
        rgb = im.convert("RGB")
        w, h = rgb.size
        if w < 400:  # upscale tiny glyphs so the VL model can read them
            scale = max(2, 400 // max(w, 1))
            rgb = rgb.resize((w * scale, h * scale), Image.LANCZOS)
        buf = io.BytesIO()
        rgb.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return blob  # best effort — never block parsing on image conversion


def _find_textbox_paragraphs(body) -> list[str]:
    """Extract text from every ``w:txbxContent`` in the document body.

    Text boxes / floating shapes are anchored to drawings but their paragraph
    content is not surfaced by ``doc.paragraphs``.
    """
    texts: list[str] = []
    for txbx in body.findall(".//w:txbxContent", _NSMAP):
        for p in txbx.findall(".//w:p", _NSMAP):
            text = _collect_text_from_element(p).strip()
            if text:
                texts.append(text)
    return texts


# ---------------------------------------------------------------------------
# Word Loader
# ---------------------------------------------------------------------------


class WordLoader(DocumentLoader):
    """Word document loader.

    Extracts paragraphs (with headings), math formulas, text-box content,
    and tables.  Uses ``lxml`` XML traversal on top of python-docx to
    capture elements the high-level API misses.
    """

    def load(self, path: str) -> Document:
        filepath = Path(path)
        if not filepath.exists():
            raise FileNotFoundError(f"文件不存在: {path}")
        if filepath.suffix.lower() not in (".docx", ".doc"):
            raise ValueError(f"不支持的文件格式: {filepath.suffix}")

        docx_doc = DocxDocument(str(filepath))

        document = Document(
            filename=filepath.name,
            file_path=str(filepath.resolve()),
            file_type="docx",
            total_pages=1,
        )

        page = Page(
            page_num=1,
            width=_DEFAULT_PAGE_WIDTH,
            height=_DEFAULT_PAGE_HEIGHT,
        )

        # ------------------------------------------------------------------
        # 1. Paragraphs (including those with embedded images and formulas)
        # ------------------------------------------------------------------
        rels = docx_doc.part.rels  # for resolving image relationships

        for order_idx, paragraph in enumerate(docx_doc.paragraphs):
            p_elem = paragraph._element  # lxml element

            # --- Build text with image placeholders at correct positions ---
            text, next_img_idx, img_positions = self._build_paragraph_text_with_images(
                p_elem, rels, page, len(page.images)
            )
            text = text.strip()

            # Detect math — if present, use math-extracted text instead
            # (OMML formulas are text-based, not image-based; they coexist
            #  with embedded images only in rare cases, so we prefer math
            #  content when available.)
            if _has_math(p_elem):
                math_text = _extract_math_latex(p_elem)
                block_type = "formula"
                text = math_text or _extract_math_text(p_elem) or _collect_text_from_element(p_elem).strip()
            else:
                # Don't skip if the paragraph only has image placeholders
                if not text and not img_positions:
                    continue
                style_name = paragraph.style.name if paragraph.style else "Normal"
                block_type = _HEADING_LEVEL_MAP.get(style_name, "text")

                # Font-size heuristic for heading classification
                font_size = self._first_font_size(paragraph)
                if block_type == "text" and font_size:
                    if font_size >= 18:
                        block_type = "title"
                    elif font_size >= 14:
                        block_type = "section_heading"

                # Content-based heading detection for test-paper structure.
                # Short text matching exam-section patterns is treated as a
                # heading so the chunker can split at section boundaries.
                if block_type == "text" and text:
                    text_stripped = text.strip()
                    if (
                        len(text_stripped) <= 60
                        and re.match(
                            r"(?:"
                            r"第[^，。\n]{0,20}[章节篇]"            # 第一章 xxx, 第2节
                            r"|(?:\d+\.)+\d*(?:\s+\S+)?"           # 1.1 xxx, 1.1.1
                            r"|[一二三四五六七八九十]+[、．]\S*"     # 一、xxx
                            r"|[（(]\d+[）)]\S*"                    # （1）xxx
                            r")",
                            text_stripped,
                        )
                    ):
                        block_type = "section_heading"

            metadata: dict = {}
            if paragraph.style:
                metadata["style"] = paragraph.style.name
            if _has_math(p_elem):
                metadata["has_formula"] = True
            if img_positions:
                metadata["embedded_image_indices"] = img_positions

            block = Block(
                content=text or "[图片]",
                block_type=block_type,
                page_num=1,
                bbox=(0, 0, _DEFAULT_PAGE_WIDTH, 20),
                reading_order=order_idx,
                confidence=1.0,
                metadata=metadata,
            )
            page.blocks.append(block)
            page.text += (text or "[图片]") + "\n"

        # ------------------------------------------------------------------
        # 2. Text-box / floating-shape content
        # ------------------------------------------------------------------
        tb_texts = _find_textbox_paragraphs(docx_doc.element.body)
        tb_offset = len(docx_doc.paragraphs)
        for i, tb_text in enumerate(tb_texts):
            block = Block(
                content=tb_text,
                block_type="text",
                page_num=1,
                bbox=(0, 0, _DEFAULT_PAGE_WIDTH, 20),
                reading_order=tb_offset + i,
                confidence=1.0,
                metadata={"source": "textbox"},
            )
            page.blocks.append(block)
            page.text += tb_text + "\n"

        # ------------------------------------------------------------------
        # 3. Tables
        # ------------------------------------------------------------------
        table_offset = tb_offset + len(tb_texts)
        for table_idx, table in enumerate(docx_doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))

            table_text = "\n".join(rows)
            if not table_text.strip():
                continue

            block = Block(
                content=table_text,
                block_type="table",
                page_num=1,
                bbox=(0, 0, _DEFAULT_PAGE_WIDTH, 20 * len(rows)),
                reading_order=table_offset + table_idx,
                confidence=1.0,
                metadata={
                    "style": "Table",
                    "num_rows": len(table.rows),
                    "num_cols": len(table.columns) if table.columns else 0,
                },
            )
            page.blocks.append(block)
            page.text += table_text + "\n"

        document.pages.append(page)
        document.status = ProcessingStatus.UPLOADED

        estimated_pages = max(1, len(page.text) // _CHARS_PER_PAGE_ESTIMATE)
        document.total_pages = estimated_pages

        return document

    # ------------------------------------------------------------------
    # Embedded image extraction (with position tracking)
    # ------------------------------------------------------------------

    def _build_paragraph_text_with_images(
        self,
        paragraph_elem,
        rels,
        page,
        start_img_idx: int,
    ) -> tuple[str, int, list[int]]:
        """Walk paragraph XML children in order, building text with image placeholders.

        In OOXML a ``w:drawing`` is nested inside ``w:r`` (a run), so we need
        to check each run's children for drawings.  When found, the image
        bytes are appended to ``page.images`` and a ``[IMG_N]`` placeholder
        is inserted at the exact text position.

        Returns
        -------
        (text_with_placeholders, next_img_idx, image_positions)
        """
        img_counter = 0
        text_parts: list[str] = []
        image_positions: list[int] = []

        try:
            for child in paragraph_elem:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

                if tag == "r":
                    # A run may contain either text (<w:t>) or a drawing (<w:drawing>)
                    drawings = child.findall(".//w:drawing", _NSMAP)
                    if drawings:
                        for drawing in drawings:
                            for blip in drawing.findall(".//a:blip", _NSMAP):
                                r_id = blip.get(
                                    f"{{{_R_NS}}}embed"
                                ) or blip.get(f"{{{_R_NS}}}link")
                                if r_id and r_id in rels:
                                    rel = rels[r_id]
                                    if "image" in (rel.reltype or ""):
                                        img_idx = start_img_idx + img_counter
                                        page.images.append(
                                            _normalize_image_bytes(rel.target_part.blob)
                                        )
                                        text_parts.append(f"[IMG_{img_idx}]")
                                        image_positions.append(img_idx)
                                        img_counter += 1
                    else:
                        for t_node in child.findall(".//w:t", _NSMAP):
                            if t_node.text:
                                text_parts.append(t_node.text)

                elif tag == "drawing":
                    # Standalone drawing (rare — some tools emit these at p level)
                    for blip in child.findall(".//a:blip", _NSMAP):
                        r_id = blip.get(f"{{{_R_NS}}}embed") or blip.get(
                            f"{{{_R_NS}}}link"
                        )
                        if r_id and r_id in rels:
                            rel = rels[r_id]
                            if "image" in (rel.reltype or ""):
                                img_idx = start_img_idx + img_counter
                                page.images.append(
                                    _normalize_image_bytes(rel.target_part.blob)
                                )
                                text_parts.append(f"[IMG_{img_idx}]")
                                image_positions.append(img_idx)
                                img_counter += 1
        except Exception:
            pass  # Best effort

        return "".join(text_parts), start_img_idx + img_counter, image_positions

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _first_font_size(paragraph) -> float | None:
        """Return the first non-``None`` font size (``pt``) in the paragraph."""
        for run in paragraph.runs:
            if run.font.size:
                return run.font.size.pt
        return None
